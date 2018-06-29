from docx import Document #python-docx을 이용하기 위해 사용한다. 워드 문서 파일 만들때 사용

document = Document()  #만들어둔 파일을 Document('여기!')에 추가를 해서 넣을 수 있습니다. 해당하는 양식이 있거나 이 이후에 추가를 해야 한다면 저 부분에 파일의 이름을 넣으면 됩니다.
#만들어둔 스타일이 있기때문에 해당하는 내용은 지워도 상관 없습니다.
document.add_heading('파이썬으로 워드파일 만들기', 0) #제목 추가
document.add_heading('첫번째로 만들어 보는 워드파일', 1)  # 0과1의 차이는 제목들의 우선순위를 나타내는 것 같다.
                                                             # 같이 첨부된 파일을 보면 알 수 있다.
document.add_heading('python-docx로 만들 수 있다.', 2)

document.add_paragraph('이 부분은 본문에 해당한다.')

#word파일 같은 경우는 자신이 스타일을 지정하거나 하는 방법으로 스타일을 저장해 둘 수가 있다. 예를들어 문단 앞에 동그라미 기호 표시를 스타일로 지정을 하고
#document.add_paragraph('이 부분은 본문에 해당한다.', style='지정한 해당 스타일 이름') -> 이런식으로 추가를 하면 해당 style로 설정이 된다.

#아래의 표도 style을 미리 지정을 하고, 불러오게 되면 설정이됩니다.
table = document.add_table(rows=1, cols=3)
#table.style = '지정한 해당 스타일 이름' -> 다음과 같습니다.
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
row_cells = table.add_row().cells
row_cells[0].text = str(123)
row_cells[1].text = str(456)
row_cells[2].text = '아무글씨'
#table을 만드는 양식이다. rows는 1행 cols는 3열

document.save('demo.docx') #docx파일을 생성한다.


#해당하는 예시는 구글에 python-docx를 검색해서 첫번째 브라우저에 들어가면 나옵니다.

#해당 예시는 가상환경의 문제로 인해서 파일을 생성하지 않았습니다.
#나중에 업무상에 필요한 내용이므로, 간단하게 알아가면 될 것 같습니다.